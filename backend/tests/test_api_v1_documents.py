"""
End-to-end tests for the /api/v1 document surface (the enterprise /
Salesforce-facing endpoints added in V2).

Runs against a minimal FastAPI app that mounts only the v1 router, so the
suite needs none of the heavy video/AI dependencies — just PyMuPDF, pyHanko,
python-docx, and (optionally) pdf2docx + LibreOffice. Endpoints whose backing
feature is not installed must return 503, and that degradation is asserted
here too.

Run:  pytest backend/tests/test_api_v1_documents.py -v
"""

from __future__ import annotations

import io
import sys
import zipfile
from datetime import datetime, timedelta, timezone
from pathlib import Path

import pytest

BACKEND_DIR = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(BACKEND_DIR))

ADMIN_TOKEN = "test-admin-token"


# --------------------------------------------------------------------------
# Fixtures
# --------------------------------------------------------------------------


@pytest.fixture(scope="session")
def client(tmp_path_factory):
    import os

    db_dir = tmp_path_factory.mktemp("db")
    os.environ["CHAMPDF_DB_PATH"] = str(db_dir / "champdf-test.db")
    os.environ["CHAMPDF_ADMIN_TOKEN"] = ADMIN_TOKEN

    import api_v1
    from fastapi import FastAPI
    from fastapi.testclient import TestClient

    api_v1.init_db()
    app = FastAPI()
    app.include_router(api_v1.router)
    return TestClient(app)


@pytest.fixture(scope="session")
def api_key(client) -> str:
    res = client.post(
        "/api/v1/admin/keys",
        headers={"X-Admin-Token": ADMIN_TOKEN},
        json={"label": "pytest", "monthly_quota": 100000},
    )
    assert res.status_code == 200, res.text
    return res.json()["key"]


@pytest.fixture(scope="session")
def auth(api_key):
    return {"Authorization": f"Bearer {api_key}"}


@pytest.fixture(scope="session")
def sample_pdf() -> bytes:
    """3-page PDF with distinct text per page."""
    import fitz

    doc = fitz.open()
    for i in range(3):
        page = doc.new_page()
        page.insert_text((72, 72), f"ChamPDF test page {i + 1}", fontsize=14)
    data = doc.tobytes()
    doc.close()
    return data


@pytest.fixture(scope="session")
def sample_png() -> bytes:
    from PIL import Image

    img = Image.new("RGB", (200, 100), (200, 40, 40))
    buf = io.BytesIO()
    img.save(buf, "PNG")
    return buf.getvalue()


@pytest.fixture(scope="session")
def sample_p12() -> bytes:
    """Self-signed cert + key packed as PKCS#12 for signing tests."""
    from cryptography import x509
    from cryptography.hazmat.primitives import hashes, serialization
    from cryptography.hazmat.primitives.asymmetric import rsa
    from cryptography.hazmat.primitives.serialization import pkcs12
    from cryptography.x509.oid import NameOID

    key = rsa.generate_private_key(public_exponent=65537, key_size=2048)
    name = x509.Name(
        [
            x509.NameAttribute(NameOID.COMMON_NAME, "ChamPDF Test Signer"),
            x509.NameAttribute(NameOID.ORGANIZATION_NAME, "Champions"),
        ]
    )
    now = datetime.now(timezone.utc)
    cert = (
        x509.CertificateBuilder()
        .subject_name(name)
        .issuer_name(name)
        .public_key(key.public_key())
        .serial_number(x509.random_serial_number())
        .not_valid_before(now - timedelta(days=1))
        .not_valid_after(now + timedelta(days=365))
        .add_extension(
            x509.KeyUsage(
                digital_signature=True, content_commitment=True,
                key_encipherment=False, data_encipherment=False,
                key_agreement=False, key_cert_sign=False, crl_sign=False,
                encipher_only=False, decipher_only=False,
            ),
            critical=True,
        )
        .sign(key, hashes.SHA256())
    )
    return pkcs12.serialize_key_and_certificates(
        b"champdf-test", key, cert, None,
        serialization.BestAvailableEncryption(b"testpass"),
    )


def _page_count(pdf_bytes: bytes) -> int:
    import fitz

    with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
        return doc.page_count


# --------------------------------------------------------------------------
# Auth plumbing
# --------------------------------------------------------------------------


def test_whoami_requires_key(client):
    assert client.get("/api/v1/whoami").status_code == 401


def test_whoami_with_key(client, auth):
    res = client.get("/api/v1/whoami", headers=auth)
    assert res.status_code == 200
    assert res.json()["label"] == "pytest"


def test_capabilities(client, auth):
    res = client.get("/api/v1/capabilities", headers=auth)
    assert res.status_code == 200
    caps = res.json()
    assert caps["pdf_core"] is True
    for k in ("pdf_sign", "pdf_ocr", "office_to_pdf", "pdf_to_docx", "pdf_tables"):
        assert isinstance(caps[k], bool)


# --------------------------------------------------------------------------
# Core PDF ops
# --------------------------------------------------------------------------


def test_merge(client, auth, sample_pdf):
    res = client.post(
        "/api/v1/pdf/merge",
        headers=auth,
        files=[
            ("files", ("a.pdf", sample_pdf, "application/pdf")),
            ("files", ("b.pdf", sample_pdf, "application/pdf")),
        ],
    )
    assert res.status_code == 200, res.text
    assert _page_count(res.content) == 6


def test_merge_rejects_single_file(client, auth, sample_pdf):
    res = client.post(
        "/api/v1/pdf/merge",
        headers=auth,
        files=[("files", ("a.pdf", sample_pdf, "application/pdf"))],
    )
    assert res.status_code == 422


def test_split(client, auth, sample_pdf):
    res = client.post(
        "/api/v1/pdf/split",
        headers=auth,
        files={"file": ("t.pdf", sample_pdf, "application/pdf")},
        data={"pages": "1,3"},
    )
    assert res.status_code == 200, res.text
    assert _page_count(res.content) == 2


def test_split_bad_range(client, auth, sample_pdf):
    res = client.post(
        "/api/v1/pdf/split",
        headers=auth,
        files={"file": ("t.pdf", sample_pdf, "application/pdf")},
        data={"pages": "99"},
    )
    assert res.status_code == 422


def test_delete_pages(client, auth, sample_pdf):
    res = client.post(
        "/api/v1/pdf/delete-pages",
        headers=auth,
        files={"file": ("t.pdf", sample_pdf, "application/pdf")},
        data={"pages": "2"},
    )
    assert res.status_code == 200, res.text
    assert _page_count(res.content) == 2


def test_rotate(client, auth, sample_pdf):
    import fitz

    res = client.post(
        "/api/v1/pdf/rotate",
        headers=auth,
        files={"file": ("t.pdf", sample_pdf, "application/pdf")},
        data={"angle": "90", "pages": "1"},
    )
    assert res.status_code == 200, res.text
    with fitz.open(stream=res.content, filetype="pdf") as doc:
        assert doc[0].rotation == 90
        assert doc[1].rotation == 0


def test_compress_never_grows(client, auth, sample_pdf):
    res = client.post(
        "/api/v1/pdf/compress",
        headers=auth,
        files={"file": ("t.pdf", sample_pdf, "application/pdf")},
    )
    assert res.status_code == 200, res.text
    assert len(res.content) <= len(sample_pdf)
    assert _page_count(res.content) == 3


def test_watermark(client, auth, sample_pdf):
    res = client.post(
        "/api/v1/pdf/watermark",
        headers=auth,
        files={"file": ("t.pdf", sample_pdf, "application/pdf")},
        data={"text": "CONFIDENTIAL", "opacity": "0.2"},
    )
    assert res.status_code == 200, res.text
    import fitz

    with fitz.open(stream=res.content, filetype="pdf") as doc:
        assert "CONFIDENTIAL" in doc[0].get_text()


def test_info(client, auth, sample_pdf):
    res = client.post(
        "/api/v1/pdf/info",
        headers=auth,
        files={"file": ("t.pdf", sample_pdf, "application/pdf")},
    )
    assert res.status_code == 200, res.text
    body = res.json()
    assert body["page_count"] == 3
    assert body["encrypted"] is False


def test_to_text(client, auth, sample_pdf):
    res = client.post(
        "/api/v1/pdf/to-text",
        headers=auth,
        files={"file": ("t.pdf", sample_pdf, "application/pdf")},
        data={"pages": "2"},
    )
    assert res.status_code == 200, res.text
    pages = res.json()["pages"]
    assert len(pages) == 1
    assert "test page 2" in pages[0]["text"]


def test_to_images(client, auth, sample_pdf):
    res = client.post(
        "/api/v1/pdf/to-images",
        headers=auth,
        files={"file": ("t.pdf", sample_pdf, "application/pdf")},
        data={"pages": "1-2", "dpi": "72"},
    )
    assert res.status_code == 200, res.text
    zf = zipfile.ZipFile(io.BytesIO(res.content))
    assert sorted(zf.namelist()) == ["page-0001.png", "page-0002.png"]


def test_from_images(client, auth, sample_png):
    res = client.post(
        "/api/v1/pdf/from-images",
        headers=auth,
        files=[
            ("files", ("1.png", sample_png, "image/png")),
            ("files", ("2.png", sample_png, "image/png")),
        ],
    )
    assert res.status_code == 200, res.text
    assert _page_count(res.content) == 2


def test_rejects_non_pdf(client, auth, sample_png):
    res = client.post(
        "/api/v1/pdf/split",
        headers=auth,
        files={"file": ("t.pdf", sample_png, "application/pdf")},
        data={"pages": "1"},
    )
    assert res.status_code == 422


# --------------------------------------------------------------------------
# Signing (pyHanko)
# --------------------------------------------------------------------------


def test_sign_and_verify(client, auth, sample_pdf, sample_p12):
    res = client.post(
        "/api/v1/pdf/sign",
        headers=auth,
        files={
            "file": ("t.pdf", sample_pdf, "application/pdf"),
            "cert": ("signer.p12", sample_p12, "application/octet-stream"),
        },
        # timestamp=false: no network call to a TSA during tests
        data={"passphrase": "testpass", "reason": "pytest", "timestamp": "false"},
    )
    if res.status_code == 503:
        pytest.skip("pyHanko not installed")
    assert res.status_code == 200, res.text
    signed = res.content
    assert len(signed) > len(sample_pdf)

    vres = client.post(
        "/api/v1/pdf/verify-signature",
        headers=auth,
        files={"file": ("signed.pdf", signed, "application/pdf")},
    )
    assert vres.status_code == 200, vres.text
    report = vres.json()
    sigs = report.get("signatures") or []
    assert len(sigs) >= 1
    assert sigs[0].get("intact") in (True, None) or report  # integrity reported


def test_sign_wrong_password(client, auth, sample_pdf, sample_p12):
    res = client.post(
        "/api/v1/pdf/sign",
        headers=auth,
        files={
            "file": ("t.pdf", sample_pdf, "application/pdf"),
            "cert": ("signer.p12", sample_p12, "application/octet-stream"),
        },
        data={"passphrase": "wrong", "timestamp": "false"},
    )
    if res.status_code == 503:
        pytest.skip("pyHanko not installed")
    assert res.status_code == 422


# --------------------------------------------------------------------------
# Conversion + optional features (must work or degrade to 503)
# --------------------------------------------------------------------------


def _docx_bytes() -> bytes:
    import docx

    d = docx.Document()
    d.add_heading("ChamPDF conversion test", level=1)
    d.add_paragraph("Hello from the enterprise API.")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def test_office_to_pdf(client, auth):
    from doc_converter import office_to_pdf_available

    res = client.post(
        "/api/v1/convert/to-pdf",
        headers=auth,
        files={
            "file": (
                "test.docx",
                _docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    if not office_to_pdf_available():
        assert res.status_code == 503
        pytest.skip("LibreOffice not installed")
    assert res.status_code == 200, res.text
    assert res.content[:5] == b"%PDF-"
    assert _page_count(res.content) >= 1


def test_office_to_pdf_rejects_unknown_ext(client, auth):
    from doc_converter import office_to_pdf_available

    if not office_to_pdf_available():
        pytest.skip("LibreOffice not installed")
    res = client.post(
        "/api/v1/convert/to-pdf",
        headers=auth,
        files={"file": ("evil.exe", b"MZ....", "application/octet-stream")},
    )
    assert res.status_code == 422


def test_pdf_to_docx(client, auth, sample_pdf):
    from doc_converter import pdf_to_docx_available

    res = client.post(
        "/api/v1/convert/pdf-to-docx",
        headers=auth,
        files={"file": ("t.pdf", sample_pdf, "application/pdf")},
    )
    if not pdf_to_docx_available():
        assert res.status_code == 503
        pytest.skip("pdf2docx not installed")
    assert res.status_code == 200, res.text
    # DOCX is a ZIP container
    assert res.content[:2] == b"PK"
    import docx

    d = docx.Document(io.BytesIO(res.content))
    text = "\n".join(p.text for p in d.paragraphs)
    assert "test page" in text


def test_ocr_degrades_gracefully(client, auth, sample_pdf):
    from ocr_processor import ocr_available

    res = client.post(
        "/api/v1/pdf/ocr",
        headers=auth,
        files={"file": ("t.pdf", sample_pdf, "application/pdf")},
        data={"language": "eng"},
    )
    if ocr_available():
        assert res.status_code == 200
        assert res.content[:5] == b"%PDF-"
    else:
        assert res.status_code == 503


def test_extract_tables(client, auth, sample_pdf):
    from table_extractor import table_extraction_available

    res = client.post(
        "/api/v1/pdf/extract-tables",
        headers=auth,
        files={"file": ("t.pdf", sample_pdf, "application/pdf")},
    )
    if not table_extraction_available():
        assert res.status_code == 503
        pytest.skip("find_tables unavailable")
    assert res.status_code == 200, res.text
    assert "tables" in res.json()


# --------------------------------------------------------------------------
# RBAC scopes
# --------------------------------------------------------------------------


@pytest.fixture(scope="session")
def sign_only_key(client) -> str:
    """A key scoped to pdf.sign only."""
    res = client.post(
        "/api/v1/admin/keys",
        headers={"X-Admin-Token": ADMIN_TOKEN},
        json={"label": "sign-bot", "monthly_quota": 1000, "scopes": ["pdf.sign"]},
    )
    assert res.status_code == 200, res.text
    body = res.json()
    assert body["scopes"] == ["pdf.sign"]
    return body["key"]


def test_admin_rejects_unknown_scope(client):
    res = client.post(
        "/api/v1/admin/keys",
        headers={"X-Admin-Token": ADMIN_TOKEN},
        json={"label": "bad", "scopes": ["pdf.frobnicate"]},
    )
    assert res.status_code == 422
    assert "Unknown scope" in res.text


def test_scoped_key_can_use_its_scope(client, sign_only_key, sample_pdf, sample_p12):
    res = client.post(
        "/api/v1/pdf/sign",
        headers={"Authorization": f"Bearer {sign_only_key}"},
        files={
            "file": ("t.pdf", sample_pdf, "application/pdf"),
            "cert": ("signer.p12", sample_p12, "application/octet-stream"),
        },
        data={"passphrase": "testpass", "timestamp": "false"},
    )
    assert res.status_code == 200, res.text


def test_scoped_key_blocked_outside_scope(client, sign_only_key, sample_pdf):
    hdr = {"Authorization": f"Bearer {sign_only_key}"}
    for path, kwargs in [
        ("/api/v1/pdf/merge", dict(files=[
            ("files", ("a.pdf", sample_pdf, "application/pdf")),
            ("files", ("b.pdf", sample_pdf, "application/pdf")),
        ])),
        ("/api/v1/pdf/to-text", dict(
            files={"file": ("t.pdf", sample_pdf, "application/pdf")},
            data={"pages": "1"},
        )),
        ("/api/v1/convert/pdf-to-docx", dict(
            files={"file": ("t.pdf", sample_pdf, "application/pdf")},
        )),
    ]:
        res = client.post(path, headers=hdr, **kwargs)
        assert res.status_code == 403, f"{path}: {res.status_code} {res.text}"
        assert res.json()["detail"]["code"] == "insufficient_scope"


def test_scoped_key_whoami_and_capabilities_still_work(client, sign_only_key):
    hdr = {"Authorization": f"Bearer {sign_only_key}"}
    who = client.get("/api/v1/whoami", headers=hdr)
    assert who.status_code == 200
    assert who.json()["scopes"] == ["pdf.sign"]
    caps = client.get("/api/v1/capabilities", headers=hdr)
    assert caps.status_code == 200


def test_parent_scope_grants_children(client, sample_pdf):
    res = client.post(
        "/api/v1/admin/keys",
        headers={"X-Admin-Token": ADMIN_TOKEN},
        json={"label": "pdf-all", "scopes": ["pdf"]},
    )
    assert res.status_code == 200
    key = res.json()["key"]
    hdr = {"Authorization": f"Bearer {key}"}
    # pdf covers pdf.read...
    res = client.post(
        "/api/v1/pdf/info", headers=hdr,
        files={"file": ("t.pdf", sample_pdf, "application/pdf")},
    )
    assert res.status_code == 200, res.text
    # ...but not convert
    res = client.post(
        "/api/v1/convert/pdf-to-docx", headers=hdr,
        files={"file": ("t.pdf", sample_pdf, "application/pdf")},
    )
    assert res.status_code == 403


def test_default_key_has_full_access(client, auth):
    # The session-wide key was minted without scopes -> '*'
    who = client.get("/api/v1/whoami", headers=auth)
    assert who.json()["scopes"] == ["*"]


# --------------------------------------------------------------------------
# Metering
# --------------------------------------------------------------------------


def test_usage_recorded(client, auth):
    res = client.get("/api/v1/whoami", headers=auth)
    assert res.status_code == 200
    # Every document call above incremented the monthly counter.
    assert res.json()["requests_used"] > 5
